import os
import glob
import sys
from docx import Document

##   DocV_Final.py

# A assinatura que identifica se o arquivo já foi alterado
SIGNATURE = "HELLO FROM DocV (Fixed)"
HEADER_MSG = f"{SIGNATURE}\nEsta é uma demonstração segura para arquivos Word.\n"

def fix_word_files():
    # Busca por todos os arquivos .docx no diretório
    for item in glob.glob("*.docx"):
        try:
            doc = Document(item)
            
            # Lê todo o texto do documento para verificar se já foi alterado
            full_text = "\n".join([p.text for p in doc.paragraphs])
            
            if SIGNATURE in full_text:
                print(f"Pulando: {item} já contém a modificação.")
                continue

            # 1. Insere o cabeçalho
            if len(doc.paragraphs) > 0:
                doc.paragraphs[0].insert_paragraph_before(HEADER_MSG)
            else:
                doc.add_paragraph(HEADER_MSG)
            
            # 2. Adiciona o conteúdo do script ao final (opcional, conforme seu original)
            with open(sys.argv[0], 'r', encoding='utf-8') as f:
                script_content = f.read()
                doc.add_section() # Adiciona uma quebra de seção para organizar
                doc.add_paragraph("\n--- SCRIPT SOURCE LOG ---\n")
                doc.add_paragraph(script_content)

            # Salva o arquivo de forma segura
            doc.save(item)
            print(f"Sucesso: {item} foi atualizado.")
            
        except Exception as e:
            print(f"Erro ao processar {item}: {e}")

if __name__ == "__main__":
    print(f"--- Iniciando processamento de arquivos .docx ---\n")
    fix_word_files()
    print(f"\n--- Tarefa concluída ---")